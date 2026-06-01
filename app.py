import streamlit as st
import fitz
import json
import pandas as pd
import plotly.express as px
import docx
import io
import os
import logging
import re
import time

from google import genai
from pydantic import BaseModel
from sklearn.ensemble import IsolationForest
from sqlalchemy import create_engine, Column, Integer, String, Float, Text, DateTime, Boolean
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import declarative_base, sessionmaker
from datetime import datetime

from auth import hash_password, verify_password
from enterprise_analytics_dashboard import render_enterprise_analytics_dashboard

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="AI Invoice Operations & Compliance Platform",
    layout="wide"
)

st.title("🛡️ AI Invoice Operations & Compliance Platform")

st.caption(
    "AI-powered invoice processing, GST validation, compliance automation, "
    "financial risk analysis, and enterprise invoice operations."
)

# =========================================================
# FOLDERS
# =========================================================

os.makedirs("logs", exist_ok=True)
os.makedirs("uploads", exist_ok=True)

# =========================================================
# LOGGING
# =========================================================

logging.basicConfig(
    filename="logs/system.log",
    level=logging.ERROR,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# =========================================================
# DATABASE
# =========================================================

DATABASE_URL = "sqlite:///invoice_platform.db"

engine = create_engine(DATABASE_URL)

SessionLocal = sessionmaker(bind=engine)

Base = declarative_base()

# =========================================================
# DATABASE MODEL
# =========================================================

class InvoiceRecord(Base):
    __tablename__ = "invoice_records"

    id = Column(Integer, primary_key=True)

    filename = Column(String)
    vendor_name = Column(String)
    invoice_number = Column(String)
    invoice_date = Column(String)

    currency = Column(String)

    subtotal = Column(Float)
    tax = Column(Float)
    total = Column(Float)

    confidence = Column(Float)
    fraud_score = Column(Float)
    compliance_score = Column(Float)

    status = Column(String)
    compliance_status = Column(String)
    approval_status = Column(String)
    risks = Column(Text)

    created_at = Column(DateTime, default=datetime.utcnow)

class User(Base):
    __tablename__ = "users"

    id = Column(Integer, primary_key=True, autoincrement=True)

    name = Column(String, nullable=False)
    email = Column(String, nullable=False, unique=True, index=True)
    password_hash = Column(String, nullable=False)

    created_at = Column(DateTime, default=datetime.utcnow)

class AuditLog(Base):
    __tablename__ = "audit_logs"

    id = Column(Integer, primary_key=True)

    action = Column(String)
    invoice_number = Column(String)

    vendor_name = Column(String)

    status = Column(String)

    timestamp = Column(
        DateTime,
        default=datetime.utcnow
    )

class NotificationRecord(Base):
    __tablename__ = "notifications"

    id = Column(Integer, primary_key=True)

    message = Column(String)

    notification_type = Column(String)

    invoice_filename = Column(String)

    created_at = Column(DateTime, default=datetime.utcnow)

    is_read = Column(Boolean, default=False)

_BCRYPT_HASH_PATTERN = re.compile(
    r"\A\$2[aby]\$\d{2}\$[./A-Za-z0-9]{53}\Z"
)

def create_user(name, email, password_hash):
    if not isinstance(name, str) or not name.strip():
        raise ValueError("Name must not be empty.")

    if not isinstance(email, str) or not email.strip():
        raise ValueError("Email must not be empty.")

    if (
        not isinstance(password_hash, str)
        or not _BCRYPT_HASH_PATTERN.fullmatch(password_hash)
    ):
        raise ValueError("A valid bcrypt password hash is required.")

    session = SessionLocal()

    try:
        user = User(
            name=name.strip(),
            email=email.strip().lower(),
            password_hash=password_hash
        )

        session.add(user)
        session.commit()
        session.refresh(user)
        session.expunge(user)

        return user

    except Exception:
        session.rollback()
        raise

    finally:
        session.close()

def get_user_by_email(email):
    if not isinstance(email, str) or not email.strip():
        return None

    session = SessionLocal()

    try:
        user = (
            session.query(User)
            .filter(User.email == email.strip().lower())
            .first()
        )

        if user:
            session.expunge(user)

        return user

    finally:
        session.close()

def initialize_auth_session():
    auth_defaults = {
        "authenticated": False,
        "user_id": None,
        "user_email": None,
        "user_name": None
    }

    for key, default_value in auth_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

def render_login_section():
    st.subheader("Login")

    with st.form("login_form"):
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        login_submitted = st.form_submit_button("Login")

    if not login_submitted:
        return

    user = get_user_by_email(email)

    if not user or not verify_password(password, user.password_hash):
        st.error("Invalid email or password")
        return

    st.session_state.authenticated = True
    st.session_state.user_id = user.id
    st.session_state.user_email = user.email
    st.session_state.user_name = user.name

    st.rerun()

def render_signup_section():
    with st.expander("Create Account"):
        with st.form("signup_form", clear_on_submit=True):
            name = st.text_input("Name")
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            confirm_password = st.text_input(
                "Confirm Password",
                type="password"
            )
            signup_submitted = st.form_submit_button("Sign Up")

        if not signup_submitted:
            return

        if not name.strip():
            st.error("Name cannot be empty.")
        elif not email.strip():
            st.error("Email cannot be empty.")
        elif not password:
            st.error("Password cannot be empty.")
        elif password != confirm_password:
            st.error("Passwords do not match.")
        elif get_user_by_email(email):
            st.error("An account with this email already exists.")
        else:
            try:
                hashed_password = hash_password(password)
                create_user(name, email, hashed_password)
            except IntegrityError:
                st.error("An account with this email already exists.")
            except ValueError as exc:
                st.error(str(exc))
            except Exception:
                logging.exception("User signup failed.")
                st.error("Unable to create account. Please try again.")
            else:
                st.success("Account created successfully. Please login.")

def render_authentication_screen():
    st.subheader("Authentication")

    login_column, signup_column = st.columns(2)

    with login_column:
        render_login_section()

    with signup_column:
        render_signup_section()

def create_notification(
    session,
    message,
    notification_type,
    invoice_filename
):
    notification = NotificationRecord(
        message=message,
        notification_type=notification_type,
        invoice_filename=invoice_filename
    )

    session.add(notification)

Base.metadata.create_all(bind=engine)

initialize_auth_session()

if not st.session_state.authenticated:
    render_authentication_screen()
    st.stop()

# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:
    st.header("⚙️ Workspace")

    try:
        gemini_api_key = st.secrets["GEMINI_API_KEY"]
        st.success("✅ AI Engine Connected")

    except Exception:
        gemini_api_key = None
        st.error("❌ Gemini API Key Missing")

    st.info("Get your key from Google AI Studio.")

    st.divider()

    st.subheader("👑 Account Tier")

    user_tier = st.selectbox(
        "Select Tier",
        [
            "Free (10 Files)",
            "Pro (1,000 Files)",
            "Enterprise (Unlimited)"
        ]
    )

    st.divider()

    user_role = st.selectbox(
        "User Role",
        [
            "Analyst",
            "Manager",
            "Auditor",
            "Admin"
        ]
    )


    if user_tier == "Free (10 Files)":
        upload_limit = 10

    elif user_tier == "Pro (1,000 Files)":
        upload_limit = 1000

    else:
        upload_limit = 10000

    st.success("Enterprise Audit Engine Online")

# =========================================================
# PYDANTIC SCHEMA
# =========================================================

class LineItem(BaseModel):
    item: str
    qty: int
    unit_price: float

class InvoiceSchema(BaseModel):
    vendor_name: str
    invoice_date: str
    invoice_number: str

    currency: str

    subtotal: float
    vat_or_tax_amount: float
    total_amount: float

    confidence_score: float

    line_items: list[LineItem]

# =========================================================
# FILE UTILITIES
# =========================================================

def save_uploaded_file(uploaded_file):
    file_path = os.path.join(
        "uploads",
        uploaded_file.name
    )

    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    return file_path

def extract_content_smart(uploaded_file):
    file_type = uploaded_file.name.split(".")[-1].lower()

    extracted_documents = []

    # =====================================================
    # PDF PROCESSING
    # =====================================================

    if file_type == "pdf":

        pdf_bytes = uploaded_file.read()

        doc = fitz.open(
            stream=pdf_bytes,
            filetype="pdf"
        )

        # PROCESS EACH PAGE AS SEPARATE INVOICE
        for page_number in range(len(doc)):

            page = doc.load_page(page_number)

            text = page.get_text()

            # ---------------------------------------------
            # SCANNED PAGE DETECTION
            # ---------------------------------------------

            if len(text.strip()) < 50:

                pix = page.get_pixmap()

                img_data = pix.tobytes("png")

                extracted_documents.append({
                    "type": "image",
                    "data": img_data,
                    "mime_type": "image/png",
                    "page": page_number + 1
                })

            else:

                extracted_documents.append({
                    "type": "text",
                    "data": text,
                    "page": page_number + 1
                })

        return extracted_documents

    # =====================================================
    # DOCX PROCESSING
    # =====================================================

    elif file_type in ["docx", "doc"]:

        document = docx.Document(uploaded_file)

        full_text = []

        for para in document.paragraphs:
            full_text.append(para.text)

        for table in document.tables:
            for row in table.rows:

                row_text = " | ".join(
                    cell.text for cell in row.cells
                )

                full_text.append(row_text)

        extracted_documents.append({
            "type": "text",
            "data": "\n".join(full_text),
            "page": 1
        })

        return extracted_documents

    # =====================================================
    # IMAGE PROCESSING
    # =====================================================

    elif file_type in ["jpg", "jpeg", "png"]:

        extracted_documents.append({
            "type": "image",
            "data": uploaded_file.getvalue(),
            "mime_type": f"image/{file_type}",
            "page": 1
        })

        return extracted_documents

    return []

# =========================================================
# AI EXTRACTION ENGINE
# =========================================================

def process_invoice_agent(api_key, content_payload):
    client = genai.Client(api_key=api_key)

    system_prompt = """
    You are an enterprise-grade AI Invoice Operations Agent.

    RULES:
    1. Extract values EXACTLY as written.
    2. Never correct mathematical mistakes.
    3. Return STRICT JSON.
    4. Lower confidence if OCR quality is weak.
    5. Detect suspicious inconsistencies.
    """

    contents = [system_prompt]

    if content_payload["type"] == "text":
        contents.append(
            f"Analyze this invoice:\n{content_payload['data']}"
        )

    else:
        contents.append("Analyze this invoice image.")

        contents.append({
            "data": content_payload["data"],
            "mime_type": content_payload["mime_type"]
        })

    for attempt in range(3):
        try:

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config={
                    "response_mime_type": "application/json",
                    "response_schema": InvoiceSchema,
                    "temperature": 0.1
                }
            )

            return json.loads(response.text)

        except Exception as e:

            if attempt < 2:
                time.sleep(2)
                continue

            raise e

# =========================================================
# FRAUD ENGINE
# =========================================================

BLACKLISTED_VENDORS = [
    "Fake Corp",
    "Shell Company Ltd",
    "Unknown Vendor"
]
CURRENCY_METADATA = {

    "INR": {
        "country": "India",
        "tax_system": "GST",
        "standard_tax_rate": 18
    },

    "USD": {
        "country": "United States",
        "tax_system": "Sales Tax",
        "standard_tax_rate": 8
    },

    "GBP": {
        "country": "United Kingdom",
        "tax_system": "VAT",
        "standard_tax_rate": 20
    },

    "EUR": {
        "country": "European Union",
        "tax_system": "VAT",
        "standard_tax_rate": 20
    },

    "AED": {
        "country": "United Arab Emirates",
        "tax_system": "VAT",
        "standard_tax_rate": 5
    },

    "AUD": {
        "country": "Australia",
        "tax_system": "GST",
        "standard_tax_rate": 10
    },

    "SGD": {
        "country": "Singapore",
        "tax_system": "GST",
        "standard_tax_rate": 9
    },

    "CAD": {
        "country": "Canada",
        "tax_system": "GST/HST",
        "standard_tax_rate": 13
    }
}

SUPPORTED_CURRENCIES = [
    "INR",
    "USD",
    "GBP",
    "EUR",
    "AED",
    "AUD",
    "SGD",
    "CAD"
]

class FraudDetector:
    def __init__(self):
        self.model = IsolationForest(
            contamination=0.05,
            random_state=42
        )

    def anomaly_score(
        self,
        historical_totals,
        current_total
    ):
        if len(historical_totals) < 5:
            return 0.1

        df = pd.DataFrame({
            "amount": historical_totals
        })

        self.model.fit(df)

        prediction = self.model.predict(
            [[current_total]]
        )

        if prediction[0] == -1:
            return 0.95

        return 0.15

def run_compliance_audit(
    data,
    fraud_history,
    historical_totals,
    existing_invoice_numbers
):
    flags = []

    currency = data.get("currency", "")

    currency_info = CURRENCY_METADATA.get(
        currency,
        {
            "country": "Unknown",
            "tax_system": "Unknown"
        }
    )

    country = currency_info["country"]
    tax_system = currency_info["tax_system"]

    standard_tax_rate = currency_info.get(
        "standard_tax_rate",
        None
    )

    if not currency:

        flags.append(
            "Currency not detected"
        )

    elif currency not in SUPPORTED_CURRENCIES:

        flags.append(
            f"Unsupported currency ({currency})"
        )

    # DUPLICATE CHECK
    invoice_id = (
        f"{data['vendor_name']}|"
        f"{data['invoice_number']}|"
        f"{data['total_amount']}"
    )

    if invoice_id in fraud_history:
        flags.append(
            f"Duplicate invoice detected ({data['invoice_number']})"
        )

    else:
        fraud_history.add(invoice_id)

    # MATH VALIDATION
    calc_subtotal = sum(
        item["qty"] * item["unit_price"]
        for item in data["line_items"]
    )

    if abs(calc_subtotal - data["subtotal"]) > 1:
        flags.append(
            "Line item subtotal mismatch detected"
        )

    calc_total = (
        data["subtotal"] +
        data["vat_or_tax_amount"]
    )

    if abs(calc_total - data["total_amount"]) > 1:
        flags.append(
            "Invoice total mismatch detected"
        )

    # DUPLICATE INVOICE CHECK
    if data["invoice_number"] in existing_invoice_numbers:
        flags.append(
            "Duplicate invoice number detected"
        )      
   # GLOBAL TAX VALIDATION

    actual_tax_rate = None

    if (
        data["subtotal"] > 0 and
        data["vat_or_tax_amount"] > 0
    ):

        actual_tax_rate = (
            data["vat_or_tax_amount"] /
            data["subtotal"]
        ) * 100

        if actual_tax_rate > 35:

            flags.append(
                f"Unusually high tax rate ({round(actual_tax_rate, 2)}%)"
            )

    if (
        actual_tax_rate is not None and
        standard_tax_rate is not None
    ):

        difference = abs(
            actual_tax_rate -
            standard_tax_rate
        )

        if difference > 20:

            flags.append(
                f"Tax rate unusually different from typical {tax_system} rate"
            )        

    # BLACKLIST CHECK
    if any(
        bad.lower() in data["vendor_name"].lower()
        for bad in BLACKLISTED_VENDORS
    ):
        flags.append(
            "Vendor appears in high-risk watchlist"
        )

    # CONFIDENCE CHECK
    if data["confidence_score"] < 0.85:
        flags.append(
            "Low confidence extraction - human review recommended"
        )

    # ML ANOMALY DETECTION

    detector = FraudDetector()
    compliance_score = 100
    compliance_score -= len(flags) * 10

    if compliance_score < 0:
        compliance_score = 0

    if compliance_score >= 90:
        compliance_status = "Compliant"

    elif compliance_score >= 70:
        compliance_status = "Review Required"

    else:
        compliance_status = "High Risk"


    fraud_score = min(
        len(flags) * 0.25,
        1.0
    )

    if fraud_score > 0.9:
        flags.append(
            "Anomalous invoice amount detected"
        )

    status = "✅ Verified"

    if compliance_score >= 90:
        compliance_status = "Compliant"
    elif compliance_score >= 70:
        compliance_status = "Review Required"
    else:
        compliance_status = "High Risk"

    if flags:
        status = "⚠️ Flagged"

    return (
        status,
        compliance_status,
        flags,
        fraud_score,
        compliance_score
    )
# =========================================================
# SESSION STATE
# =========================================================

if "fraud_history" not in st.session_state:
    st.session_state["fraud_history"] = set()

if "historical_totals" not in st.session_state:
    st.session_state["historical_totals"] = []

# =========================================================
# TABS
# =========================================================
session = SessionLocal()

unread_count = (
    session.query(NotificationRecord)
    .filter(
        NotificationRecord.is_read == False
    )
    .count()
)

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📤 Invoice Processing & Compliance",
    "📊 Analytics Dashboard",
    "🗄️ Invoice Database",
    "⚖️ Audit Trail",
    f"📢 Notifications ({unread_count})"
])

# =========================================================
# TAB 1
# =========================================================

with tab1:
    st.subheader("📤 Upload Source Documents")

    uploaded_files = st.file_uploader(
        f"Supported: PDF, PNG, JPG, DOCX (Limit: {upload_limit})",
        type=[
            "pdf",
            "docx",
            "png",
            "jpg",
            "jpeg"
        ],
        accept_multiple_files=True
    )

    if uploaded_files and gemini_api_key:
        if len(uploaded_files) > upload_limit:
            st.error("Tier limit exceeded.")

        else:
            if st.button(
                "🚀 Start Invoice Operations Pipeline"
            ):
                with st.spinner(
                    "Processing invoices..."
                ):
                    batch_data = []

                    progress_bar = st.progress(0)

                    session = SessionLocal()

                    for idx, file in enumerate(uploaded_files):
                        try:
                            # SAVE FILE
                            save_uploaded_file(file)

                            # EXTRACT CONTENT
                            payloads = extract_content_smart(file)

                            if not payloads:
                                continue

                            # PROCESS EACH PAGE SEPARATELY
                            for payload in payloads:

                                try:
                                    data = process_invoice_agent(
                                        gemini_api_key,
                                        payload
                                    )

                                except Exception as e:

                                    logging.error(str(e))

                                    st.error(
                                        "AI service temporarily unavailable. Please try again later."
                                    )

                                    continue

                                # AUDIT ENGINE
                                existing_invoice_numbers = [
                                    r.invoice_number
                                    for r in session.query(
                                        InvoiceRecord
                                    ).all()
                                ]
                                status, compliance_status, risks, fraud_score, compliance_score = run_compliance_audit(
                                    data,
                                    st.session_state["fraud_history"],
                                    st.session_state["historical_totals"],
                                    existing_invoice_numbers
                                )

                                if compliance_status == "High Risk":

                                    create_notification(
                                        session,
                                        f"High risk invoice detected: {data['invoice_number']}",
                                        "HIGH_RISK",
                                        file.name
                                    )

                                st.session_state[
                                    "historical_totals"
                                ].append(
                                    data["total_amount"]
                                )

                                # DATABASE SAVE
                                record = InvoiceRecord(
                                    filename=f"{file.name} | Page {payload['page']}",
                                    vendor_name=data["vendor_name"],
                                    invoice_number=data["invoice_number"],
                                    invoice_date=data["invoice_date"],
                                    currency=data["currency"],
                                    subtotal=data["subtotal"],
                                    tax=data["vat_or_tax_amount"],
                                    total=data["total_amount"],
                                    confidence=data["confidence_score"],
                                    fraud_score=fraud_score,
                                    compliance_score=compliance_score,
                                    status=status,
                                    compliance_status=compliance_status,
                                    approval_status="Pending Review",
                                    risks=", ".join(risks)
                                )

                                session.add(record)
                                create_notification(
                                    session,
                                    f"Invoice processed: {data['invoice_number']}",
                                    "PROCESSING",
                                    file.name
                                )
                                audit = AuditLog(
                                    action="Invoice Processed",
                                    invoice_number=data["invoice_number"],
                                    vendor_name=data["vendor_name"],
                                    status=status
                                )

                                session.add(audit)
                                session.commit()

                                # UI DATA
                                batch_data.append({
                                    "Filename": f"{file.name} | Page {payload['page']}",
                                    "Vendor": data["vendor_name"],
                                    "Invoice #": data["invoice_number"],
                                    "Currency": data["currency"],
                                    "Total": data["total_amount"],
                                    "Fraud Score": round(fraud_score, 2),
                                    "Compliance Score": round(compliance_score, 2),
                                    "Confidence": f"{data['confidence_score']*100:.0f}%",
                                    "Status": status,
                                    "Approval Status": "Pending Review",
                                    "Risks": ", ".join(risks)
                                })

                        except Exception as e:
                            logging.error(str(e))
                            continue

                        progress_bar.progress(
                            (idx + 1) / len(uploaded_files)
                        )

                    # RESULTS
                    if batch_data:
                        st.success(
                            "Invoice processing completed."
                        )

                        df = pd.DataFrame(batch_data)

                        st.subheader(
                            "📂 Enterprise Audit Ledger"
                        )

                        st.dataframe(
                            df,
                            use_container_width=True
                        )

                        # METRICS
                        st.divider()

                        col1, col2, col3 = st.columns(3)

                        flagged_count = len(
                            df[df["Status"] == "⚠️ Flagged"]
                        )

                        avg_confidence = round(
                            df["Confidence"]
                            .str.replace("%", "")
                            .astype(float)
                            .mean(),
                            2
                        )

                        col1.metric(
                            "Processed Invoices",
                            len(df)
                        )

                        col2.metric(
                            "Flagged Invoices",
                            flagged_count
                        )

                        col3.metric(
                            "Average Confidence",
                            f"{avg_confidence}%"
                        )

                        # EXPORTS
                        st.divider()

                        st.subheader(
                            "💾 Export Reports"
                        )

                        ec1, ec2, ec3 = st.columns(3)

                        # CSV
                        csv = df.to_csv(
                            index=False
                        ).encode("utf-8")

                        ec1.download_button(
                            "📄 Download CSV",
                            csv,
                            "invoice_audit.csv",
                            "text/csv",
                            use_container_width=True
                        )

                        # JSON
                        json_data = df.to_json(
                            orient="records"
                        ).encode("utf-8")

                        ec2.download_button(
                            "🤖 Download JSON",
                            json_data,
                            "invoice_audit.json",
                            "application/json",
                            use_container_width=True
                        )

                        # EXCEL
                        output = io.BytesIO()

                        with pd.ExcelWriter(
                            output,
                            engine="openpyxl"
                        ) as writer:
                            df.to_excel(
                                writer,
                                index=False
                            )

                        ec3.download_button(
                            "📊 Download Excel",
                            output.getvalue(),
                            "invoice_audit.xlsx",
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )

# =========================================================
# TAB 2
# =========================================================

with tab2:

    session = SessionLocal()

    rows = session.query(
        InvoiceRecord
    ).all()

    analytics_records = []

    for row in rows:
        analytics_records.append({
            "vendor_name": row.vendor_name,
            "invoice_number": row.invoice_number,
            "invoice_date": row.invoice_date,
            "currency": row.currency,
            "total_amount": row.total,
            "fraud_score": row.fraud_score * 100,
            "status": row.status,
            "compliance_status": row.compliance_status
        })

    analytics_df = pd.DataFrame(
        analytics_records
    )

    if not analytics_df.empty:
        render_enterprise_analytics_dashboard(
            analytics_df,
            gemini_model=None
        )

        fig = px.bar(
            analytics_df,
            x="vendor_name",
            y="total_amount",
            color="status",
            title="Vendor Spend Analytics"
        )

        st.plotly_chart(
            fig,
            use_container_width=True
        )

        fig2 = px.histogram(
            analytics_df,
            x="fraud_score",
            nbins=10,
            title="Fraud Score Distribution"
        )

        st.plotly_chart(
            fig2,
            use_container_width=True
        )

    else:
        st.info(
            "Upload invoices to start generating compliance analytics."
        )

# =========================================================
# TAB 3
# =========================================================

with tab3:
    st.subheader("🗄️ Stored Invoice Database")

    session = SessionLocal()

    rows = session.query(
        InvoiceRecord
    ).all()

    database_records = []

    for row in rows:
        database_records.append({
            "Filename": row.filename,
            "Vendor": row.vendor_name,
            "Invoice": row.invoice_number,
            "Date": row.invoice_date,
            "Currency": row.currency,
            "Total": row.total,
            "Fraud Score": row.fraud_score,
            "Status": row.status,
            "Approval Status": row.approval_status,
            "Risks": row.risks
        })

    if database_records:

        db_df = pd.DataFrame(
            database_records
        )

        st.dataframe(
            db_df,
            use_container_width=True
        )

        st.divider()

        st.subheader(
            "📝 Invoice Approval Center"
        )

        pending_rows = [
            r for r in rows
            if r.approval_status == "Pending Review"
        ]

        if pending_rows:

            invoice_options = [
                f"{r.invoice_number} | {r.vendor_name}"
                for r in pending_rows
            ]

            selected_invoice = st.selectbox(
                "Select Invoice",
                invoice_options
            )

            if user_role in ["Manager", "Admin"]:
                col1, col2 = st.columns(2)

                with col1:
                    approve_btn = st.button(
                        "✅ Approve Invoice"
                    )

                with col2:
                    reject_btn = st.button(
                        "❌ Reject Invoice"
                    )

                if approve_btn:

                    selected_invoice_number = (
                        selected_invoice.split("|")[0].strip()
                    )

                    invoice_record = (
                        session.query(InvoiceRecord)
                        .filter(
                            InvoiceRecord.invoice_number
                            == selected_invoice_number
                        )
                        .first()
                    )

                    if invoice_record:

                        invoice_record.approval_status = "Approved"

                        create_notification(
                            session,
                            f"Invoice approved: {invoice_record.invoice_number}",
                            "APPROVAL",
                            invoice_record.filename
                        )

                        audit_entry = AuditLog(
                            action="Invoice Approved",
                            invoice_number=invoice_record.invoice_number,
                            vendor_name=invoice_record.vendor_name,
                            status="Approved"
                        )

                        session.add(
                            audit_entry
                        )

                        session.commit()

                        st.success(
                            "Invoice approved successfully."
                        )

                        st.rerun()

                    else:
                        st.error(
                            "Selected invoice not found."
                        )

                if reject_btn:

                    selected_invoice_number = (
                        selected_invoice.split("|")[0].strip()
                    )
                    invoice_record = (
                        session.query(InvoiceRecord)
                        .filter(
                            InvoiceRecord.invoice_number
                            == selected_invoice_number
                        )
                        .first()
                    )

                    if invoice_record:
                        invoice_record.approval_status = "Rejected"

                        create_notification(
                            session,
                            f"Invoice rejected: {invoice_record.invoice_number}",
                            "REJECTION",
                            invoice_record.filename
                        )
                        audit_entry = AuditLog(
                            action="Invoice Rejected",
                            invoice_number=invoice_record.invoice_number,
                            vendor_name=invoice_record.vendor_name,
                            status="Rejected"
                        )

                        session.add(
                            audit_entry
                        )

                        session.commit()

                        st.success(
                            "Invoice rejected successfully."
                        )

                        st.rerun()

                    else:
                        st.error(
                            "Selected invoice not found."
                        )

            else:
                st.warning(
                    "Only Managers and Admins can approve invoices."
                )

        else:
            st.success(
                "No invoices pending review."
            )
    else:

        st.info(
            "No invoices stored yet."
        )
# =========================================================
# TAB 4
# =========================================================

with tab4:

    if user_role not in ["Manager", "Auditor", "Admin"]:

        st.warning(
            "Audit Trail access is restricted to Managers, Auditors and Admins."
        )

    else:

        st.subheader("📜 Audit Trail")

        audit_status_filter = st.selectbox(
            "Filter by Status",
            ["All", "Verified", "Flagged"]
        )

        audit_search = st.text_input(
            "Search Invoice Number or Vendor"
        )

        session = SessionLocal()

        audit_records = (
            session.query(AuditLog)
            .order_by(AuditLog.timestamp.desc())
            .all()
        )

        if audit_records:

            audit_df = pd.DataFrame([
                {
                    "Action": r.action,
                    "Invoice #": r.invoice_number,
                    "Vendor": r.vendor_name,
                    "Status": r.status,
                    "Timestamp": r.timestamp
                }
                for r in audit_records
            ])

            filtered_audit_df = audit_df.copy()

            if audit_status_filter != "All":

                filtered_audit_df = filtered_audit_df[
                    filtered_audit_df["Status"].str.contains(
                        audit_status_filter,
                        case=False,
                        na=False
                    )
                ]

            if audit_search:

                filtered_audit_df = filtered_audit_df[
                    filtered_audit_df["Invoice #"]
                    .astype(str)
                    .str.contains(
                        audit_search,
                        case=False,
                        na=False
                    )
                    |
                    filtered_audit_df["Vendor"]
                    .astype(str)
                    .str.contains(
                        audit_search,
                        case=False,
                        na=False
                    )
                ]

            st.dataframe(
                filtered_audit_df,
                use_container_width=True
            )

        else:

            st.info(
                "No audit records available."
            )

# =========================================================
# TAB 5
# =========================================================

with tab5:

    st.subheader("📢 Notification Center")

    col1, col2 = st.columns(2)

    with col1:
        mark_all_read = st.button(
            "✅ Mark All Read"
        )

    with col2:
        refresh_notifications = st.button(
            "🔄 Refresh"
        )

    session = SessionLocal()

    if mark_all_read:

        session.query(
            NotificationRecord
        ).update(
            {
                NotificationRecord.is_read: True
            }
        )

        session.commit()

        st.success(
            "All notifications marked as read."
        )

        st.rerun()

    notifications = (
        session.query(NotificationRecord)
        .order_by(
            NotificationRecord.created_at.desc()
        )
        .all()
    )

    if notifications:

        notification_data = []

        for n in notifications:

            notification_data.append({
                "Type": n.notification_type,
                "Message": n.message,
                "Invoice": n.invoice_filename,
                "Created": n.created_at,
                "Status": (
                    "🔵 Unread"
                    if not n.is_read
                    else "✅ Read"
                )
            })

        notification_df = pd.DataFrame(
            notification_data
        )

        st.dataframe(
            notification_df,
            use_container_width=True
        )

    else:

        st.info(
            "No notifications available."
        )

# =========================================================
# FOOTER
# =========================================================

st.markdown("---")

st.caption(
    "Built with AI-powered invoice intelligence and compliance automation."
)
